#!/usr/bin/env python3
"""Linux eVoice sync: docs/ → audios/*.mp3 via Piper or espeak-ng + ffmpeg.

Usage:
  python3 linux_sync.py --project-dir /var/tmp/evoice-jobs/<id>/project [--only name] [--premium]

Emits frequent progress lines and a final STATS line.
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
import tempfile
import urllib.error
import urllib.request
from pathlib import Path

DOC_EXTENSIONS = {
    ".docx",
    ".txt",
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".tif",
    ".tiff",
    ".bmp",
    ".gif",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp", ".gif"}

PREMIUM_SYSTEM = (
    "Eres un editor de guiones hablados en español. Reescribe el texto para "
    "síntesis de voz (TTS): oraciones cortas y claras, expande abreviaturas, "
    "quita markdown/URLs ruidosas, mantén el significado. Responde SOLO con el "
    "texto hablado, sin títulos ni explicaciones."
)


def log(msg: str) -> None:
    print(msg, flush=True)


def needs_regen(doc: Path, mp3: Path) -> bool:
    if not mp3.is_file():
        return True
    return doc.stat().st_mtime > mp3.stat().st_mtime


def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(". ".join(cells))
    return "\n\n".join(parts)


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def extract_image(path: Path) -> str:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageOps

    log(f"EXTRACT {path.name} pct=10 detail=open_image")
    image = Image.open(path).convert("RGB")
    w, h = image.size
    if max(w, h) < 2500:
        image = image.resize((w * 2, h * 2), Image.Resampling.LANCZOS)
    log(f"EXTRACT {path.name} pct=40 detail=preprocess")
    gray = ImageOps.autocontrast(ImageOps.grayscale(image))
    gray = ImageEnhance.Contrast(gray).enhance(1.4)
    log(f"EXTRACT {path.name} pct=70 detail=tesseract")
    text = pytesseract.image_to_string(gray, lang="spa+eng").strip()
    log(f"EXTRACT {path.name} pct=100 detail=chars={len(text)}")
    return text


def load_doc_text(path: Path) -> str:
    ext = path.suffix.lower()
    log(f"EXTRACT {path.name} pct=5 detail=start")
    if ext == ".txt":
        text = extract_txt(path)
        log(f"EXTRACT {path.name} pct=100 detail=chars={len(text)}")
        return text
    if ext == ".docx":
        text = extract_docx(path)
        log(f"EXTRACT {path.name} pct=100 detail=chars={len(text)}")
        return text
    if ext == ".pdf":
        text = extract_pdf(path)
        log(f"EXTRACT {path.name} pct=100 detail=chars={len(text)}")
        return text
    if ext in IMAGE_EXTENSIONS:
        return extract_image(path)
    raise ValueError(f"unsupported extension: {ext}")


def premium_optimize(text: str, name: str) -> str:
    """DeepSeek chat completions with stream=true; emit PREMIUM progress as SSE chunks arrive."""
    key = (os.environ.get("DEEPSEEK_API_KEY") or "").strip()
    if not key:
        raise RuntimeError("DEEPSEEK_API_KEY not configured for premium")
    model = (os.environ.get("DEEPSEEK_MODEL_REASONING") or "deepseek-v4-pro").strip()
    base = (os.environ.get("DEEPSEEK_BASE_URL") or "https://api.deepseek.com").rstrip("/")
    log(f"PREMIUM {name} pct=5 detail=deepseek_stream_start model={model}")
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": PREMIUM_SYSTEM},
            {"role": "user", "content": text[:120000]},
        ],
        "temperature": 0.3,
        "stream": True,
    }
    req = urllib.request.Request(
        f"{base}/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {key}",
            "Accept": "text/event-stream",
        },
        method="POST",
    )
    parts: list[str] = []
    last_pct = 5
    try:
        with urllib.request.urlopen(req, timeout=300) as resp:
            while True:
                raw = resp.readline()
                if not raw:
                    break
                line = raw.decode("utf-8", errors="replace").strip()
                if not line:
                    continue
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if data == "[DONE]":
                    break
                try:
                    chunk = json.loads(data)
                except json.JSONDecodeError:
                    continue
                delta = (
                    chunk.get("choices", [{}])[0]
                    .get("delta", {})
                    .get("content", "")
                )
                if not delta:
                    continue
                parts.append(delta)
                n = sum(len(p) for p in parts)
                # Map growing output length into 10–95 so the UI moves before completion.
                pct = min(95, 10 + (n // 40))
                if pct >= last_pct + 5 or (pct > last_pct and pct >= 90):
                    last_pct = pct
                    snippet = delta.replace("\n", " ")[:80]
                    log(f"PREMIUM {name} pct={pct} detail=chars={n} {snippet}")
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")[:400]
        raise RuntimeError(f"deepseek HTTP {exc.code}: {detail}") from exc
    content = "".join(parts).strip()
    if not content:
        raise RuntimeError("deepseek returned empty content")
    log(f"PREMIUM {name} pct=100 detail=optimized {len(text)}→{len(content)} chars")
    return content


def find_ffmpeg() -> str:
    which = shutil.which("ffmpeg")
    if not which:
        raise FileNotFoundError("ffmpeg not found on PATH")
    return which


def text_to_wav_piper(text: str, wav_path: Path) -> None:
    piper = shutil.which("piper")
    if not piper:
        raise FileNotFoundError("piper not found")
    model = Path(__file__).resolve().parent / "models" / "es_ES-sharvard-medium.onnx"
    env_model = Path(os.environ.get("EVOICE_PIPER_MODEL", "")).expanduser()
    if env_model.is_file():
        model = env_model
    if not model.is_file():
        raise FileNotFoundError(f"piper model missing: {model}")
    proc = subprocess.run(
        [piper, "--model", str(model), "--output_file", str(wav_path)],
        input=text.encode("utf-8"),
        capture_output=True,
        check=False,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.decode("utf-8", errors="replace") or "piper failed")


def text_to_wav_espeak(text: str, wav_path: Path) -> None:
    espeak = shutil.which("espeak-ng") or shutil.which("espeak")
    if not espeak:
        raise FileNotFoundError("espeak-ng not found")
    proc = subprocess.run(
        [espeak, "-v", "es", "-w", str(wav_path), text],
        capture_output=True,
        check=False,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.decode("utf-8", errors="replace") or "espeak failed")


def wav_to_mp3(wav_path: Path, mp3_path: Path) -> None:
    ffmpeg = find_ffmpeg()
    log(f"FFMPEG {mp3_path.name} pct=50 detail=encoding")
    proc = subprocess.run(
        [
            ffmpeg,
            "-y",
            "-i",
            str(wav_path),
            "-codec:a",
            "libmp3lame",
            "-qscale:a",
            "4",
            str(mp3_path),
        ],
        capture_output=True,
        check=False,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.decode("utf-8", errors="replace")[-400:] or "ffmpeg failed")
    log(f"FFMPEG {mp3_path.name} pct=100 detail=ok")


def chunk_text(text: str, max_chars: int = 800) -> list[str]:
    text = text.strip()
    if not text:
        return []
    parts: list[str] = []
    buf: list[str] = []
    size = 0
    for para in text.replace("\r", "").split("\n"):
        para = para.strip()
        if not para:
            continue
        if size + len(para) + 1 > max_chars and buf:
            parts.append(" ".join(buf))
            buf = [para]
            size = len(para)
        else:
            buf.append(para)
            size += len(para) + 1
    if buf:
        parts.append(" ".join(buf))
    return parts or [text]


def text_to_mp3(text: str, mp3_path: Path, name: str, tmp_parent: Path | None = None) -> None:
    text = text.strip()
    if not text:
        raise ValueError("empty text")
    parent = tmp_parent if tmp_parent is not None else mp3_path.parent
    parent.mkdir(parents=True, exist_ok=True)
    chunks = chunk_text(text)
    with tempfile.TemporaryDirectory(prefix="evoice-tts-", dir=str(parent)) as tmp:
        tmp_path = Path(tmp)
        wav_parts: list[Path] = []
        for i, chunk in enumerate(chunks, start=1):
            pct = int(100 * (i - 1) / max(len(chunks), 1))
            log(f"TTS {name} pct={pct} detail=chunk {i}/{len(chunks)}")
            wav = tmp_path / f"part-{i:04d}.wav"
            try:
                text_to_wav_piper(chunk, wav)
                if i == 1:
                    log(f"TTS {name} detail=engine=piper")
            except Exception as piper_err:  # noqa: BLE001
                if i == 1:
                    log(f"TTS {name} detail=piper_fail ({piper_err}); espeak-ng")
                text_to_wav_espeak(chunk, wav)
            wav_parts.append(wav)
            log(f"TTS {name} pct={int(100 * i / max(len(chunks), 1))} detail=chunk_done")
        if len(wav_parts) == 1:
            wav_to_mp3(wav_parts[0], mp3_path)
            return
        list_file = tmp_path / "concat.txt"
        list_file.write_text(
            "\n".join(f"file '{p.resolve().as_posix()}'" for p in wav_parts),
            encoding="utf-8",
        )
        merged = tmp_path / "merged.wav"
        ffmpeg = find_ffmpeg()
        proc = subprocess.run(
            [ffmpeg, "-y", "-f", "concat", "-safe", "0", "-i", str(list_file), "-c", "copy", str(merged)],
            capture_output=True,
            check=False,
        )
        if proc.returncode != 0:
            raise RuntimeError(proc.stderr.decode("utf-8", errors="replace")[-400:] or "ffmpeg concat failed")
        wav_to_mp3(merged, mp3_path)


def sync_project(
    project_dir: Path,
    only: set[str] | None = None,
    *,
    premium: bool = False,
) -> dict[str, int]:
    docs_dir = project_dir / "docs"
    audios_dir = project_dir / "audios"
    audios_dir.mkdir(parents=True, exist_ok=True)
    docs = sorted(
        p
        for p in docs_dir.iterdir()
        if p.is_file()
        and p.name != ".keep"
        and not p.name.endswith(".premium.txt")
        and p.suffix.lower() in DOC_EXTENSIONS
        and (not only or p.name in only)
    )
    stats = {"docs": len(docs), "generated": 0, "skipped": 0, "failed": 0}
    log(f"STEP convert docs={len(docs)} premium={int(premium)}")
    if not docs:
        log("No convertible files in docs/")
        return stats
    for idx, doc in enumerate(docs, start=1):
        mp3 = audios_dir / f"{doc.stem}.mp3"
        log(f"FILE {doc.name} state=active")
        log(f"STEP convert doc={idx}/{len(docs)} file={doc.name}")
        if not needs_regen(doc, mp3):
            log(f"skip  {doc.name} (mp3 up to date)")
            log(f"FILE {doc.name} state=skipped")
            stats["skipped"] += 1
            continue
        reason = "missing mp3" if not mp3.is_file() else "doc newer than mp3"
        log(f"gen   {doc.name} -> audios/{mp3.name} ({reason})")
        try:
            text = load_doc_text(doc)
            if not text.strip():
                raise ValueError("No readable text extracted")
            if premium:
                text = premium_optimize(text, doc.name)
                premium_path = docs_dir / f"{doc.stem}.premium.txt"
                premium_path.write_text(text, encoding="utf-8")
                log(f"PREMIUM {doc.name} detail=wrote {premium_path.name}")
            text_to_mp3(text, mp3, doc.name, tmp_parent=project_dir)
            stats["generated"] += 1
            log(f"ok     {doc.name} -> {mp3.name}")
            log(f"FILE {doc.name} state=done")
        except Exception as exc:  # noqa: BLE001
            stats["failed"] += 1
            log(f"FAIL  {doc.name}: {exc}")
            log(f"FILE {doc.name} state=failed")
    return stats


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-dir", type=Path, required=True)
    parser.add_argument(
        "--only",
        action="append",
        default=[],
        help="Only convert this basename (repeatable). Default: all convertible docs.",
    )
    parser.add_argument(
        "--premium",
        action="store_true",
        help="Optimize extracted text with DeepSeek reasoning before TTS.",
    )
    args = parser.parse_args()
    project_dir = args.project_dir.resolve()
    if not (project_dir / "docs").is_dir():
        log(f"docs/ missing under {project_dir}")
        return 1
    only = {n for n in args.only if n} or None
    stats = sync_project(project_dir, only=only, premium=bool(args.premium))
    log(
        f"STATS docs={stats['docs']} generated={stats['generated']} "
        f"skipped={stats['skipped']} failed={stats['failed']}"
    )
    return 0 if stats["failed"] == 0 else 2


if __name__ == "__main__":
    sys.exit(main())
